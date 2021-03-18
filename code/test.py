# -*- coding: utf-8 -*-
"""
Created on Tue Mar  9 08:16:29 2021

@author: YU Yixiong
"""

import static as st
import Heli
import City
import Mission
import matplotlib.pyplot as plt
import docx

def AddHeadText(text, size):
    title_ = document.add_heading(level=3)
    # title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER# 标题居中
    title_run = title_.add_run(text)  # 添加标题内容
    title_run.font.size = Pt(size)  # 设置标题字体大小
    title_run.font.name = 'Times New Roman'  # 设置标题西文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # 设置标题中文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)#字体颜色

#添加段落内容(参数1：文本内容，参数2：字体大小，参数3：上行距,参数4：字体粗细，参数5：段落位置)
def AddParaText(text, size, space, thickness, position):
    p = document.add_paragraph()  # 段落
    #判断居中还是靠左,0为靠左
    if position == 0:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT #靠左
    else :
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #居中
    p.paragraph_format.space_before = Pt(space)
    text = p.add_run(text)
    #判断字体是否加粗（1为不加粗）
    if thickness == 1:
        text.bold = False
    else:
        text.bold = True #加粗
    text.font.name = 'Times New Roman'
    text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    text.font.size = Pt(size)

def createReport():
    doc=docx.Document(r"..\\data\\template.docx")
    doc.add_paragraph("jkl")
    doc.save("..\\data\\fleet1\\planShow.docx")

createReport()