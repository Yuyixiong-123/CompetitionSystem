# -*- coding: utf-8 -*-
"""
Created on Thu Mar 18 17:07:36 2021

@author: YU Yixiong
"""


import json
def getLog(pathheliLog,pathMissionLog):
    with open(pathheliLog,"r") as f:
        heliLog=json.load(f)
        f.close()
    with open(pathMissionLog,"r") as f:
        mLog=json.load(f)
        f.close()
    
    return heliLog,mLog
heliLog,mLog=getLog("../data/fleet1/HeliLog.json", "../data/fleet1/MissionLog.json")

def getStr(l):
    s=''
    for e in l:
       s+=str(e) 
    return s
    
import docx
document=docx.Document(r"..\data\template.docx")

document.add_heading('推演过程', 1)

table = document.add_table(rows=len(heliLog), cols=2, style='Table Grid')
rowList=['']*len(heliLog)
for i in range(len(heliLog)):
    # rowList[i]=table.rows[i].cells
    # rowList[i][0].txt=heliLog[i][-1]/60
    # rowList[i][1].txt=getStr(heliLog[i])
    print(heliLog[i][-1]/60)
    table.cell(i,0).text=str(heliLog[i][-1]/60)
    table.cell(i,1).text=getStr(heliLog[i])
    print(i)
# table.cell(1,0).txt=heliLog[1][-1]/60   
# document.add_heading('推演过程', 1)


    
document.save(r"..\data\fleet1\planShow.docx")